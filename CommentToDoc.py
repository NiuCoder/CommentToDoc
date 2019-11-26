
# coding: utf-8


import re
from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
import os


# 构建正则表达式
fileNameStr = '文件名:\s+(\w+\.h)'
descStr = '描\s?述:\s+(\w+模块)'
patStr = '\*\s\@name\s+([^\n]+)\n\*\s\@brief\s+([^\n]+)\n((?:\*\s\@param\[in\]\s+[^\n]+\n)*)((?:\*\s\@param\[out\]\s+[^\n]+\n)*)' +'\*\s\@return\s+([^\n]+)\n\*\s\@note\s+([^\n]+)\n'+'.*\/\n'+'(\w+\s\w+(?:[^;]*\n?);)'
inputStr = '\@param\[in\]\s+([^\n]+)\n'
outputStr = '\@param\[out\]\s+([^\n]+)\n'
conStr = '([\w\*]+)'


def iter_heading(paragraphs):
    for paragraph in paragraphs:
        isItHeading=re.match('Heading ([1-9])',paragraph.style.name)
        if isItHeading:
            yield int(isItHeading.groups()[0]),paragraph

def addHeaderNumbering(document):
    hNums=[0,0,0,0,0]
    for index,hx in iter_heading(document.paragraphs):
        # ---put zeroes below---
        for i in range(index+1,5):
            hNums[i]=0
        # ---increment this---
        hNums[index]+=1
        # ---prepare the string---
        hStr=""
        for i in range(1,index+1):
            hStr+="%d."%hNums[i]
        # ---add the numbering---
        hx.text=hStr+" "+hx.text


# 遍历dicList写入doc
def writeToDoc(desc,filename,modelname,dicList):
    document = Document()
    document.add_heading(modelname+'软件接口文档',0)
    document.add_heading('', level = 1).add_run(modelname)
    para_desc = document.add_paragraph(desc)
    para_desc.paragraph_format.left_indent = Inches(0.5)
    for dic in dicList:
        document.add_heading('', level = 2).add_run(dic['name'])
        document.add_paragraph('功能概述', style='List Number')
        para_brief = document.add_paragraph(dic['brief'])
        para_brief.paragraph_format.left_indent = Inches(0.5)
        document.add_paragraph('函数原型', style='List Number')
        para_pro = document.add_paragraph(dic['protype'])
        para_pro.paragraph_format.left_indent = Inches(0.5)
        document.add_paragraph('输入', style='List Number')
        inrowNum = len(dic['input'])
        if inrowNum > 0:
            table = document.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '序号'
            hdr_cells[1].text = '名称'
            hdr_cells[2].text = '类型'
            hdr_cells[3].text = '含义'
            hdr_cells[4].text = '说明'
            for index,(ty,nam,dec) in enumerate(dic['input']):
                row_cells = table.add_row().cells
                row_cells[0].text = str(index+1)
                row_cells[1].text = nam
                row_cells[2].text = ty
                row_cells[3].text = dec
        else:
            para_input = document.add_paragraph('无')
            para_input.paragraph_format.left_indent = Inches(0.5)
        document.add_paragraph('输出', style='List Number')
        outrowNum = len(dic['output'])
        if outrowNum > 0:
            table = document.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '序号'
            hdr_cells[1].text = '名称'
            hdr_cells[2].text = '类型'
            hdr_cells[3].text = '含义'
            hdr_cells[4].text = '说明'
            for index,(ty,nam,dec) in enumerate(dic['output']):
                row_cells = table.add_row().cells
                row_cells[0].text = str(index+1)
                row_cells[1].text = nam
                row_cells[2].text = ty
                row_cells[3].text = dec
        else:
            para_output = document.add_paragraph('无')
            para_output.paragraph_format.left_indent = Inches(0.5)
        document.add_paragraph('其他', style='List Number')
        para_note = document.add_paragraph(dic['note'])
        para_note.paragraph_format.left_indent = Inches(0.5)
    addHeaderNumbering(document)
    document.save(modelname+'软件接口文档'+'.docx')


def patMatch(content):
    dicList = []
    # 一级匹配
    res = re.finditer(patStr, content)
    # 填充字典
    for mat in res:
        # 字典中的元素分别是name,brief,input,output,return,note,protype
        comm = {}
        comm['name'] = mat.group(1)
        comm['brief'] = mat.group(2)
        inputRes = mat.group(3)
        comm['input'] = []
        if len(inputRes) != 0:
            res2 = re.finditer(inputStr, inputRes)
            for mat2 in res2:
                if len(mat2.group(1)) != 0:
                    res3 = re.finditer(conStr, mat2.group(1))
                    inputDict = []
                    for mat3 in res3:
                        if mat3.group(1) != '无':
                            inputDict.append(mat3.group(1))
                    if len(inputDict) > 0:
                        comm['input'].append(inputDict)
        outputRes = mat.group(4)
        comm['output'] = []
        if len(outputRes) != 0:
            res4 = re.finditer(outputStr, outputRes)
            for mat4 in res4:
                if len(mat4.group(1)) != 0:
                    res5 = re.finditer(conStr, mat4.group(1))
                    outputDict = []
                    for mat5 in res5:
                        if mat5.group(1) != '无':
                            outputDict.append(mat5.group(1))
                    if len(outputDict) > 0:
                        comm['output'].append(outputDict)
        comm['return'] = mat.group(5)
        comm['note'] = mat.group(6)
        comm['protype'] = mat.group(7)
        dicList.append(comm)
    return dicList


# In[15]:

def commentToDoc():
    # 读取文件
    path = os.getcwd()
    for file in os.listdir(path):
        if not os.path.isdir(file) and os.path.splitext(file)[-1] == '.h':
            fo = open(file,encoding='gbk',errors='ignore')
            content = ''
            for line in fo.readlines():
                content += line
            modelname = re.search(descStr, content).group(1)
            desc = modelname+"接口头文件:"+file
            dictList = []
            dictList = patMatch(content)
            if len(dictList) > 0:
                writeToDoc(desc,file,modelname,dictList)


if __name__ == '__main__':
	commentToDoc()



